import PyPDF2
from docx import Document
import os
from typing import Tuple

class ResumeParser:
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
        return text

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
        return text

    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extract text from resume file (PDF or DOCX)"""
        _, file_extension = os.path.splitext(file_path)
        file_extension = file_extension.lower()
        
        if file_extension == ".pdf":
            return ResumeParser.extract_text_from_pdf(file_path)
        elif file_extension == ".docx":
            return ResumeParser.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
